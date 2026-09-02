# -*- coding: utf-8 -*-
"""
Turn whatever a client sends us (chat text, email, .docx, .pdf, .xlsx)
into a single plain-text string that the extractor can work on.

Design note: we deliberately keep this dumb and format-specific rather than
trying to build one universal parser. Each format has a tiny, well-tested
function; the pipeline picks one based on the file extension.
"""
from __future__ import annotations

import email
from email import policy
from pathlib import Path

import docx
import openpyxl
import pdfplumber


def read_txt(path: str | Path) -> str:
    return Path(path).read_text(encoding="utf-8", errors="replace")


def read_eml(path: str | Path) -> str:
    """Very small .eml reader: pulls subject + plain-text body."""
    with open(path, "rb") as f:
        msg = email.message_from_binary_file(f, policy=policy.default)
    parts = []
    subject = msg.get("subject")
    if subject:
        parts.append(f"Тема: {subject}")
    body = msg.get_body(preferencelist=("plain", "html"))
    if body is not None:
        parts.append(body.get_content())
    return "\n".join(parts)


def read_docx(path: str | Path) -> str:
    d = docx.Document(str(path))
    chunks = [p.text for p in d.paragraphs if p.text.strip()]
    for table in d.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                chunks.append(" | ".join(cells))
    return "\n".join(chunks)


def read_pdf(path: str | Path) -> str:
    chunks = []
    with pdfplumber.open(str(path)) as pdf:
        for page in pdf.pages:
            text = page.extract_text() or ""
            if text.strip():
                chunks.append(text)
            for table in page.extract_tables() or []:
                for row in table:
                    cells = [str(c).strip() for c in row if c and str(c).strip()]
                    if cells:
                        chunks.append(" | ".join(cells))
    return "\n".join(chunks)


def read_xlsx(path: str | Path) -> str:
    wb = openpyxl.load_workbook(str(path), data_only=True)
    chunks = []
    for ws in wb.worksheets:
        for row in ws.iter_rows(values_only=True):
            cells = [str(c).strip() for c in row if c is not None and str(c).strip()]
            if cells:
                chunks.append(" | ".join(cells))
    return "\n".join(chunks)


_READERS = {
    ".txt": read_txt,
    ".eml": read_eml,
    ".docx": read_docx,
    ".pdf": read_pdf,
    ".xlsx": read_xlsx,
    ".xlsm": read_xlsx,
}


def extract_text(path: str | Path) -> str:
    """Dispatch on file extension. Raises ValueError for unsupported types."""
    p = Path(path)
    ext = p.suffix.lower()
    if ext not in _READERS:
        raise ValueError(
            f"Unsupported file type '{ext}'. Supported: {sorted(_READERS)}"
        )
    return _READERS[ext](p).strip()
